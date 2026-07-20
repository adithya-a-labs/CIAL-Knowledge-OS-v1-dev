from __future__ import annotations
from datetime import datetime, timezone
import hashlib
import json
from pathlib import Path
from types import SimpleNamespace
import uuid
import zipfile
import pytest
from backend.app.services.export_document import ExportDocument, ExportSource, MarkdownExportParser, cited_reference_ids
from backend.app.services.export_document import BulletListBlock, NumberedListBlock, ParagraphBlock
from backend.app.services.export_renderers import BODY_LEADING_PT, BODY_PT, HEADING_PT, TITLE_PT, DocxRenderer, HtmlPreviewRenderer, PdfRenderer, PdfTaskItem, PdfUnsupportedGlyphError, VectorCheckbox, _pdf_text, _plain
from backend.app.services.export_service import ExportService, sanitize_filename

MARKDOWN = """# Operations
Use **bold** and *careful* steps [1].

- Inspect the light
- Record the result

1. Notify operations
2. Validate reopening

| Check | Result |
|---|---|
| Light | Ready |

```text
status = ready
```

> Preserve the cited limitation [2].

https://example.com/a/very/long/path
"""

def document():
    parser=MarkdownExportParser(); return ExportDocument("Runway response","Professional export",datetime.now(timezone.utc),"What should we do?",{},parser.parse(MARKDOWN),cited_reference_ids(MARKDOWN),[ExportSource(1,"Manual.pdf",7),ExportSource(2,"Policy.docx",3)],{})

def test_markdown_it_builds_canonical_blocks_and_citations():
    value=document(); names={type(block).__name__ for block in value.blocks}
    assert {"HeadingBlock","ParagraphBlock","BulletListBlock","NumberedListBlock","TableBlock","CodeBlock","QuoteBlock"}.issubset(names)
    assert value.citations==[1,2]

def test_pdf_and_docx_are_independently_valid(tmp_path:Path):
    value=document(); pdf=tmp_path/"out.pdf"; docx=tmp_path/"out.docx"
    PdfRenderer().render(value,pdf); DocxRenderer().render(value,docx)
    assert pdf.read_bytes().startswith(b"%PDF") and pdf.stat().st_size>100
    assert zipfile.is_zipfile(docx)
    with zipfile.ZipFile(docx) as archive: assert {"[Content_Types].xml","word/document.xml"}.issubset(archive.namelist())

def test_docx_preview_is_sanitized_and_has_no_active_content():
    value=document(); value.blocks.extend(MarkdownExportParser().parse('<script>alert(1)</script><img src=x onerror=alert(1)>[x](javascript:alert(1))'))
    html=HtmlPreviewRenderer().render(value).casefold()
    assert "<script" not in html and "onerror" not in html and "javascript:" not in html
    assert "content-security-policy" not in html  # CSP is an HTTP response header

def test_filename_is_backend_sanitized():
    name=sanitize_filename('../bad:<title>. ',"pdf",datetime(2026,7,20,10,30,tzinfo=timezone.utc))
    assert name=="CIAL-Knowledge-OS_bad-title_2026-07-20_10-30.pdf"
    assert "/" not in name and "\\" not in name

def test_genuinely_unsupported_text_fails_with_codepoint_detail(tmp_path:Path):
    value=ExportDocument("Airport ✈",None,datetime.now(timezone.utc),None,{},[],[],[],{})
    with pytest.raises(PdfUnsupportedGlyphError,match=r"unsupported_glyph: U\+2708"): PdfRenderer().render(value,tmp_path/"unicode.pdf")

def task_document():
    markdown="""## Action Checklist

- [ ] Ensure the runway light is inspected before reopening [3][7]
- [x] Record the completed inspection [3]
- Ordinary supporting note

1. Notify operations
2. Validate the result

Literal [ ] text in a normal paragraph remains literal.

### Immediate

- [ ] Perform the immediate action [7]
"""
    parser=MarkdownExportParser();return ExportDocument("Action checklist",None,datetime.now(timezone.utc),None,{},parser.parse(markdown),[3,7],[],{})

def test_task_items_are_normalized_once_and_other_markdown_is_preserved():
    value=task_document();bullet=next(block for block in value.blocks if isinstance(block,BulletListBlock))
    assert bullet.task_states==[False,True,None]
    assert "[ ]" not in _plain(bullet.items[0][0].inline_nodes)
    assert "[x]" not in _plain(bullet.items[1][0].inline_nodes)
    assert _plain(bullet.items[2][0].inline_nodes)=="Ordinary supporting note"
    assert any(isinstance(block,NumberedListBlock) for block in value.blocks)
    paragraph=next(block for block in value.blocks if isinstance(block,ParagraphBlock) and "Literal" in _plain(block.inline_nodes))
    assert "[ ]" in _plain(paragraph.inline_nodes)

def test_html_preview_has_single_noninteractive_task_markers_and_readable_citations():
    html=HtmlPreviewRenderer().render(task_document())
    assert html.count("☐")==2 and html.count("☒")==1
    assert "• [ ]" not in html and "• [x]" not in html
    assert '<ul class="task-list">' in html and '<span class="task-marker">' in html
    assert '<p class="meta">' in html
    assert "[3]</span> <span class=\"citation\">[7]" in html
    lowered=html.casefold();assert "<input" not in lowered and "<script" not in lowered and "onerror" not in lowered and "javascript:" not in lowered and " style=" not in lowered

def test_docx_tasks_use_hanging_indent_without_word_bullet_numbering(tmp_path:Path):
    from docx import Document
    path=tmp_path/"tasks.docx";DocxRenderer().render(task_document(),path);doc=Document(path)
    task_paragraphs=[p for p in doc.paragraphs if p.text.startswith(("☐","☒"))]
    assert len(task_paragraphs)==3
    for paragraph in task_paragraphs:
        assert paragraph.style.name=="Normal"
        assert paragraph.paragraph_format.left_indent is not None
        assert paragraph.paragraph_format.first_line_indent is not None and paragraph.paragraph_format.first_line_indent.pt<0
    assert not any(p.text.startswith("• ☐") or p.text.startswith("• ☒") for p in doc.paragraphs)

def test_pdf_tasks_are_vector_only_and_keep_other_list_types(tmp_path:Path):
    from pypdf import PdfReader
    path=tmp_path/"tasks.pdf";PdfRenderer().render(task_document(),path);text="\n".join(page.extract_text() or "" for page in PdfReader(path).pages)
    assert not any(marker in text for marker in ("☐","☒","✓","✔","• [ ]","• [x]"))
    assert "Ordinary supporting note" in text
    assert "1. Notify operations" in text and "2. Validate the result" in text
    assert "[3] [7]" in text and "[ ] text in a normal paragraph" in text

def test_vector_checkbox_draws_square_and_checked_mark():
    class CanvasSpy:
        def __init__(self):self.rects=[];self.lines=[]
        def saveState(self):pass
        def restoreState(self):pass
        def setStrokeColorRGB(self,*value):pass
        def setLineWidth(self,value):pass
        def rect(self,*args,**kwargs):self.rects.append((args,kwargs))
        def line(self,*args):self.lines.append(args)
    unchecked=VectorCheckbox(False);unchecked.canv=CanvasSpy();unchecked.draw()
    checked=VectorCheckbox(True);checked.canv=CanvasSpy();checked.draw()
    assert len(unchecked.canv.rects)==1 and unchecked.canv.lines==[]
    assert len(checked.canv.rects)==1 and len(checked.canv.lines)==2

def test_pdf_task_flowable_has_hanging_column_and_splits_without_duplicate_marker():
    from reportlab.lib.styles import getSampleStyleSheet
    item=PdfTaskItem("wrapped "*1200,getSampleStyleSheet()["BodyText"],False)
    parts=item.split(300,100)
    assert item.marker_width>item.paragraph.style.fontSize
    assert len(parts)>1 and parts[0].show_marker and all(not part.show_marker for part in parts[1:])

def test_many_pdf_tasks_paginate_and_use_builtin_fonts(monkeypatch,tmp_path:Path):
    from pypdf import PdfReader
    from reportlab.pdfbase import pdfmetrics
    monkeypatch.setattr(pdfmetrics,"registerFont",lambda *args,**kwargs: (_ for _ in ()).throw(AssertionError("font registration is not allowed")))
    markdown="## Action Checklist\n\n"+"\n".join(f"- [{'x' if index%2 else ' '}] Task item {index} with enough text to wrap across the available line width [3]" for index in range(90))
    value=ExportDocument("Checklist",None,datetime.now(timezone.utc),None,{},MarkdownExportParser().parse(markdown),[3],[],{})
    path=tmp_path/"many-tasks.pdf";PdfRenderer().render(value,path);reader=PdfReader(path);text="\n".join(page.extract_text() or "" for page in reader.pages)
    assert len(reader.pages)>1 and "Task item 0" in text and "Task item 89" in text
    assert not any(marker in text for marker in ("☐","☒","✓","✔","[x]"))

def test_pdf_smart_punctuation_uses_bounded_ascii_fallback(tmp_path:Path):
    from pypdf import PdfReader
    value=ExportDocument("“Runway”—status…",None,datetime.now(timezone.utc),None,{},[],[],[],{})
    path=tmp_path/"punctuation.pdf";PdfRenderer().render(value,path);text="\n".join(page.extract_text() or "" for page in PdfReader(path).pages)
    assert _pdf_text("‘one’ “two” – three — four…","test")=="'one' \"two\" - three -- four..."
    assert _pdf_text("Café Œuvre","test")=="Café Œuvre"
    assert '"Runway"--status...' in text

def test_docx_keeps_unicode_checkbox_and_non_latin_text(tmp_path:Path):
    from dataclasses import replace
    from docx import Document
    value=replace(task_document(),title="Checklist മലയാളം");path=tmp_path/"unicode.docx";DocxRenderer().render(value,path)
    text="\n".join(paragraph.text for paragraph in Document(path).paragraphs)
    assert "Checklist മലയാളം" in text and "☐" in text and "☒" in text

def _service_job(content: str):
    snapshot={"content":content,"citations":[],"sources":[],"metadata":{},"query":None,"message_created_at":datetime.now(timezone.utc).isoformat()}
    canonical=json.dumps(snapshot,sort_keys=True,separators=(",",":"),ensure_ascii=False)
    return SimpleNamespace(id=uuid.uuid4(),user_id=uuid.uuid4(),message_id=uuid.uuid4(),format="pdf",status="queued",progress_stage="queued",progress_percent=0,title="Action checklist",options={"include_sources":True,"include_generated_timestamp":True},source_snapshot=snapshot,source_content_hash=hashlib.sha256(canonical.encode()).hexdigest(),started_at=None,completed_at=None,storage_key=None,preview_storage_key=None,output_filename=None,output_mime_type=None,file_size_bytes=None,error_code=None,safe_error_message=None)

def _install_service_db(monkeypatch,job):
    import backend.app.services.export_service as module
    class Db:
        def refresh(self,value):pass
        def commit(self):pass
    db=Db()
    class Context:
        def __enter__(self):return db
        def __exit__(self,*args):pass
    monkeypatch.setattr(module,"SessionLocal",lambda:Context())
    monkeypatch.setattr(module,"ExportRepository",lambda value:SimpleNamespace(get=lambda export_id:job))

def test_successful_pdf_job_reaches_ready_and_artifact_is_available(monkeypatch,tmp_path:Path):
    import backend.app.api.routes.exports as routes
    job=_service_job("## Action Checklist\n\n- [ ] Inspect [3]\n- [x] Record [3]");_install_service_db(monkeypatch,job)
    service=ExportService(tmp_path);service.process(job.id)
    artifact=service.artifact(job)
    assert job.status=="ready" and job.progress_stage=="ready" and job.progress_percent==100
    assert artifact.read_bytes().startswith(b"%PDF") and service.artifact(job,preview=False)==artifact
    monkeypatch.setattr(routes,"_owned",lambda *args:job)
    request=SimpleNamespace(app=SimpleNamespace(state=SimpleNamespace(export_service=service)));db=SimpleNamespace(commit=lambda:None)
    preview=routes.preview_export(job.id,request,db);download=routes.download_export(job.id,request,db)
    assert Path(preview.path)==artifact and preview.media_type=="application/pdf"
    assert Path(download.path)==artifact and download.media_type=="application/pdf"

def test_failed_pdf_job_cleans_partial_file_and_uses_safe_error(monkeypatch,tmp_path:Path):
    job=_service_job("Unsupported ✈");_install_service_db(monkeypatch,job)
    service=ExportService(tmp_path);service.process(job.id)
    assert job.status=="failed" and job.progress_stage=="failed" and job.error_code=="unsupported_glyph"
    assert job.safe_error_message=="unsupported_glyph: U+2708 is not supported by the packaged PDF fonts"
    assert not list(tmp_path.rglob("*.tmp")) and job.storage_key is None

def test_bounded_typography_constants_and_compact_preview_page():
    assert TITLE_PT==24 and BODY_PT==11 and BODY_LEADING_PT==16
    assert HEADING_PT=={1:18,2:14,3:12}
    html=HtmlPreviewRenderer().render(task_document())
    assert "max-width:760px" in html and "font-size:32px" in html and "font:16px/1.52" in html
