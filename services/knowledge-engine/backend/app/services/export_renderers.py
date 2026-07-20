"""Independent PDF, DOCX, and sanitized HTML renderers for ExportDocument."""
from __future__ import annotations
from html import escape
import logging
from pathlib import Path
import re
import bleach
from reportlab.platypus import Flowable, Paragraph
from backend.app.services.export_document import *

logger = logging.getLogger(__name__)

TITLE_PT = 24
BODY_PT = 11
BODY_LEADING_PT = 16
HEADING_PT = {1: 18, 2: 14, 3: 12}
HEADING_SPACE_BEFORE_PT = {1: 12, 2: 10, 3: 8}
HEADING_SPACE_AFTER_PT = {1: 5, 2: 4, 3: 3}

def _plain(nodes) -> str:
    parts = []
    previous_citation = False
    for node in nodes:
        if isinstance(node, (InlineText, InlineCode)): parts.append(node.text); previous_citation = False
        elif isinstance(node, CitationReference):
            separator = " " if previous_citation else ("\u00a0" if parts and not parts[-1].endswith((" ","\n","\u00a0")) else "")
            parts.append(f"{separator}[{node.reference_id}]"); previous_citation = True
        elif isinstance(node, SoftBreak): parts.append("\n"); previous_citation = False
        elif hasattr(node, "children"): parts.append(_plain(node.children)); previous_citation = False
    return "".join(parts)

_PDF_PUNCTUATION_FALLBACKS = str.maketrans({
    "\u00a0": " ", "\u2018": "'", "\u2019": "'", "\u201c": '"', "\u201d": '"',
    "\u2013": "-", "\u2014": "--", "\u2026": "...",
})

class PdfUnsupportedGlyphError(ValueError):
    code = "unsupported_glyph"

def _pdf_text(value: str, block_kind: str) -> str:
    """Apply bounded WinAnsi-safe punctuation fallbacks without deleting text."""
    normalized = value.translate(_PDF_PUNCTUATION_FALLBACKS)
    unsupported = None
    for character in normalized:
        try:
            character.encode("cp1252")
        except UnicodeEncodeError:
            unsupported = character
            break
    if unsupported is not None:
        logger.warning(
            "pdf_unsupported_glyph",
            extra={"code_point": f"U+{ord(unsupported):04X}", "font": "Helvetica", "block_kind": block_kind},
        )
        raise PdfUnsupportedGlyphError(
            f"unsupported_glyph: U+{ord(unsupported):04X} is not supported by the packaged PDF fonts"
        )
    return normalized

def _pdf_inline(nodes, block_kind: str = "paragraph") -> str:
    parts = []
    previous_citation = False
    for node in nodes:
        if isinstance(node, InlineText): parts.append(escape(_pdf_text(node.text, block_kind)))
        elif isinstance(node, InlineCode): parts.append(f'<font name="Courier">{escape(_pdf_text(node.text, block_kind))}</font>')
        elif isinstance(node, CitationReference):
            separator = " " if previous_citation else ("&#160;" if parts else "")
            parts.append(f"{separator}[{node.reference_id}]"); previous_citation = True; continue
        elif isinstance(node, SoftBreak): parts.append("<br/>")
        elif isinstance(node, Bold): parts.append(f"<b>{_pdf_inline(node.children, block_kind)}</b>")
        elif isinstance(node, Italic): parts.append(f"<i>{_pdf_inline(node.children, block_kind)}</i>")
        elif isinstance(node, Link): parts.append(f'<a href="{escape(node.href, quote=True)}">{_pdf_inline(node.children, block_kind)}</a>')
        previous_citation = False
    return "".join(parts)

class VectorCheckbox(Flowable):
    """Small, font-independent checkbox drawn entirely with PDF vectors."""
    def __init__(self, checked: bool, size: float = 8.5) -> None:
        super().__init__(); self.checked=checked; self.box_size=size; self.width=size; self.height=BODY_LEADING_PT
    def wrap(self, avail_width, avail_height): return self.width, self.height
    def draw(self) -> None:
        y=self.height-self.box_size-2
        self.canv.saveState(); self.canv.setStrokeColorRGB(0.14,0.36,0.13); self.canv.setLineWidth(0.8)
        self.canv.rect(0,y,self.box_size,self.box_size,stroke=1,fill=0)
        if self.checked:
            self.canv.line(1.5,y+4.0,3.5,y+1.8); self.canv.line(3.5,y+1.8,self.box_size-1.2,y+self.box_size-1.4)
        self.canv.restoreState()

class PdfTaskItem(Flowable):
    """Splittable task text with one vector marker and a stable hanging column."""
    marker_width = 18
    def __init__(self, content: str, style, checked: bool, show_marker: bool = True, paragraph=None) -> None:
        super().__init__(); self.paragraph=paragraph or Paragraph(content,style); self.checked=checked; self.show_marker=show_marker; self.spaceAfter=3
    def wrap(self, avail_width, avail_height):
        _,height=self.paragraph.wrap(max(1,avail_width-self.marker_width),avail_height); self.width=avail_width; self.height=height; return avail_width,height
    def split(self, avail_width, avail_height):
        parts=self.paragraph.split(max(1,avail_width-self.marker_width),avail_height)
        if len(parts)<=1:return []
        return [PdfTaskItem("",part.style,self.checked,self.show_marker and index==0,paragraph=part) for index,part in enumerate(parts)]
    def draw(self) -> None:
        self.paragraph.drawOn(self.canv,self.marker_width,0)
        if self.show_marker:VectorCheckbox(self.checked).drawOn(self.canv,0,max(0,self.height-BODY_LEADING_PT))

class PdfRenderer:
    mime_type = "application/pdf"
    def render(self, document: ExportDocument, output: Path) -> None:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import SimpleDocTemplate, Spacer, Table, TableStyle, Preformatted, HRFlowable, PageBreak
        styles = getSampleStyleSheet(); body = ParagraphStyle("CIALBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=BODY_PT, leading=BODY_LEADING_PT, spaceAfter=5)
        title_style=ParagraphStyle("CIALTitle",parent=styles["Title"],fontName="Helvetica-Bold",fontSize=TITLE_PT,leading=29,textColor=colors.HexColor("#245c20"),spaceAfter=5)
        heading = {n: ParagraphStyle(f"H{n}", parent=styles[f"Heading{min(n,3)}"], fontName="Helvetica-Bold",fontSize=HEADING_PT[n],leading=HEADING_PT[n]+4,spaceBefore=HEADING_SPACE_BEFORE_PT[n],spaceAfter=HEADING_SPACE_AFTER_PT[n],textColor=colors.HexColor("#245c20"), keepWithNext=True) for n in (1,2,3)}
        story = [Paragraph(escape(_pdf_text(document.title,"title")), title_style)]
        if document.context_metadata.get("export_options",{}).get("include_generated_timestamp",True): story.append(Paragraph(f"Generated {document.generated_at:%d %b %Y, %H:%M UTC}", styles["Normal"]))
        story.append(Spacer(1, 8*mm))
        if document.query: story += [Paragraph("Question", heading[2]), Paragraph(escape(_pdf_text(document.query,"query")), body)]
        for block in document.blocks:
            if isinstance(block, HeadingBlock): story.append(Paragraph(escape(_pdf_text(block.text,"heading")), heading[min(block.level,3)]))
            elif isinstance(block, ParagraphBlock): story.append(Paragraph(_pdf_inline(block.inline_nodes), body))
            elif isinstance(block, (BulletListBlock, NumberedListBlock)):
                for idx, item in enumerate(block.items, start=getattr(block,"start",1)):
                    text = " ".join(_pdf_inline(part.inline_nodes,"list_item") for part in item if isinstance(part, ParagraphBlock))
                    task_state = block.task_states[idx-1] if isinstance(block,BulletListBlock) and idx-1 < len(block.task_states) else None
                    if task_state is not None:
                        story.append(PdfTaskItem(text,body,task_state))
                    else:
                        prefix = "•" if isinstance(block, BulletListBlock) else f"{idx}."; story.append(Paragraph(f"{prefix} {text}", ParagraphStyle("list", parent=body, leftIndent=12,firstLineIndent=-8,spaceAfter=3)))
            elif isinstance(block, TableBlock):
                data = [[Paragraph(escape(_pdf_text(cell,"table")), body) for cell in block.headers]] + [[Paragraph(escape(_pdf_text(cell,"table")), body) for cell in row] for row in block.rows]
                if data: 
                    table = Table(data, repeatRows=1, hAlign="LEFT"); table.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#e8f1e5")),("GRID",(0,0),(-1,-1),.4,colors.HexColor("#b8c8b4")),("VALIGN",(0,0),(-1,-1),"TOP"),("PADDING",(0,0),(-1,-1),5)])); story.append(table); story.append(Spacer(1,6))
            elif isinstance(block, CodeBlock): story.append(Preformatted(_pdf_text(block.code,"code"), ParagraphStyle("code", fontName="Courier", fontSize=8, leading=10, backColor=colors.HexColor("#f3f5f2"), borderPadding=6)))
            elif isinstance(block, QuoteBlock):
                for part in block.blocks:
                    if isinstance(part, ParagraphBlock): story.append(Paragraph(_pdf_inline(part.inline_nodes), ParagraphStyle("quote", parent=body, leftIndent=12, borderColor=colors.HexColor("#7ca475"), borderWidth=1, borderPadding=6)))
            elif isinstance(block, HorizontalRuleBlock): story.append(HRFlowable(color=colors.HexColor("#b8c8b4")))
            elif isinstance(block, PageBreakBlock): story.append(PageBreak())
        if document.sources:
            story += [PageBreak(), Paragraph("Sources", heading[1])]
            for source in document.sources:
                location = f", p. {source.page_number}" if source.page_number else (f", {source.location}" if source.location else "")
                story.append(Paragraph(f"[{source.citation_number}] {escape(_pdf_text(source.document_title,'source'))}{escape(_pdf_text(location,'source'))}", body))
        def footer(canvas, doc):
            canvas.saveState(); canvas.setFont("Helvetica", 8); canvas.setFillColor(colors.HexColor("#526050")); canvas.drawString(20*mm, 12*mm, "CIAL Knowledge OS"); canvas.drawRightString(190*mm, 12*mm, f"Page {doc.page}"); canvas.restoreState()
        SimpleDocTemplate(str(output), pagesize=A4, rightMargin=20*mm, leftMargin=20*mm, topMargin=22*mm, bottomMargin=20*mm, title=document.title, author="CIAL Knowledge OS").build(story, onFirstPage=footer, onLaterPages=footer)

class DocxRenderer:
    mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    def render(self, document: ExportDocument, output: Path) -> None:
        from docx import Document
        from docx.enum.section import WD_SECTION
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Mm, Pt
        doc = Document(); section = doc.sections[0]; section.page_width=Mm(210); section.page_height=Mm(297); section.left_margin=section.right_margin=Mm(20); section.top_margin=section.bottom_margin=Mm(22)
        normal=doc.styles["Normal"]; normal.font.name="Arial"; normal.font.size=Pt(BODY_PT); normal.paragraph_format.space_after=Pt(4);normal.paragraph_format.line_spacing=1.4
        title=doc.styles["Title"];title.font.name="Arial";title.font.size=Pt(TITLE_PT);title.font.bold=True;title.paragraph_format.space_after=Pt(5)
        for level in (1,2,3):
            style=doc.styles[f"Heading {level}"];style.font.name="Arial";style.font.size=Pt(HEADING_PT[level]);style.font.bold=True;style.paragraph_format.space_before=Pt(HEADING_SPACE_BEFORE_PT[level]);style.paragraph_format.space_after=Pt(HEADING_SPACE_AFTER_PT[level]);style.paragraph_format.keep_with_next=True
        doc.core_properties.title=document.title; doc.core_properties.subject="AI Assistant response export"; doc.core_properties.author="CIAL Knowledge OS"; doc.core_properties.created=document.generated_at
        doc.add_heading(document.title, 0)
        if document.context_metadata.get("export_options",{}).get("include_generated_timestamp",True): doc.add_paragraph(f"Generated {document.generated_at:%d %b %Y, %H:%M UTC}")
        if document.query: doc.add_heading("Question", 2); doc.add_paragraph(document.query)
        for block in document.blocks:
            if isinstance(block, HeadingBlock): doc.add_heading(block.text, min(block.level,3))
            elif isinstance(block, ParagraphBlock): self._add_inline(doc.add_paragraph(),block.inline_nodes)
            elif isinstance(block, (BulletListBlock, NumberedListBlock)):
                for index,item in enumerate(block.items):
                    text=" ".join(_plain(p.inline_nodes) for p in item if isinstance(p,ParagraphBlock));state=block.task_states[index] if isinstance(block,BulletListBlock) and index<len(block.task_states) else None
                    if state is None: p=doc.add_paragraph(text,style="List Bullet" if isinstance(block,BulletListBlock) else "List Number");p.paragraph_format.space_after=Pt(3)
                    else:
                        p=doc.add_paragraph();p.paragraph_format.left_indent=Mm(9);p.paragraph_format.first_line_indent=Mm(-7);p.paragraph_format.space_after=Pt(3);p.paragraph_format.line_spacing=1.3;p.add_run("☒" if state else "☐");p.add_run("\t"+text)
            elif isinstance(block, TableBlock):
                table=doc.add_table(rows=1, cols=max(1,len(block.headers))); table.style="Table Grid"
                for i,value in enumerate(block.headers): table.rows[0].cells[i].text=value
                for row in block.rows:
                    cells=table.add_row().cells
                    for i,value in enumerate(row[:len(cells)]): cells[i].text=value
            elif isinstance(block, CodeBlock): p=doc.add_paragraph(); run=p.add_run(block.code); run.font.name="Courier New"; run.font.size=Pt(8)
            elif isinstance(block, QuoteBlock):
                for part in block.blocks:
                    if isinstance(part, ParagraphBlock): doc.add_paragraph(_plain(part.inline_nodes), style="Intense Quote")
            elif isinstance(block, HorizontalRuleBlock): doc.add_paragraph("─"*48)
            elif isinstance(block, PageBreakBlock): doc.add_page_break()
        if document.sources:
            doc.add_page_break(); doc.add_heading("Sources",1)
            for source in document.sources: doc.add_paragraph(f"[{source.citation_number}] {source.document_title}" + (f", p. {source.page_number}" if source.page_number else ""))
        header=section.header.paragraphs[0]; header.text="CIAL Knowledge OS"; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
        footer=section.footer.paragraphs[0]; footer.text="CIAL Knowledge OS • Confidential knowledge export"; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
        doc.save(output)

    def _add_inline(self,paragraph,nodes) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        previous_citation=False
        for node in nodes:
            if isinstance(node, InlineText): paragraph.add_run(node.text);previous_citation=False
            elif isinstance(node, CitationReference): paragraph.add_run((" " if previous_citation else "\u00a0")+f"[{node.reference_id}]");previous_citation=True
            elif isinstance(node, SoftBreak): paragraph.add_run().add_break();previous_citation=False
            elif isinstance(node, InlineCode): run=paragraph.add_run(node.text);run.font.name="Courier New";previous_citation=False
            elif isinstance(node,(Bold,Italic)):
                start=len(paragraph.runs);self._add_inline(paragraph,node.children)
                for run in paragraph.runs[start:]: run.bold=isinstance(node,Bold);run.italic=isinstance(node,Italic)
                previous_citation=False
            elif isinstance(node,Link):
                relationship=paragraph.part.relate_to(node.href,"http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",is_external=True)
                hyperlink=OxmlElement("w:hyperlink");hyperlink.set(qn("r:id"),relationship);run=OxmlElement("w:r");properties=OxmlElement("w:rPr");color=OxmlElement("w:color");color.set(qn("w:val"),"245C20");properties.append(color);run.append(properties);text=OxmlElement("w:t");text.text=_plain(node.children);run.append(text);hyperlink.append(run);paragraph._p.append(hyperlink);previous_citation=False

class HtmlPreviewRenderer:
    def render(self, document: ExportDocument) -> str:
        def safe_text(value: str) -> str:
            value = re.sub(r"<script\b[^>]*>.*?</script\s*>", "", value, flags=re.IGNORECASE | re.DOTALL)
            value = re.sub(r"\bon[a-z]+\s*=\s*[^\s>]+", "", value, flags=re.IGNORECASE)
            value = re.sub(r"javascript\s*:", "", value, flags=re.IGNORECASE)
            return escape(value)
        def inline_html(nodes) -> str:
            parts=[];previous_citation=False
            for node in nodes:
                if isinstance(node,InlineText):parts.append(safe_text(node.text));previous_citation=False
                elif isinstance(node,InlineCode):parts.append(f"<code>{safe_text(node.text)}</code>");previous_citation=False
                elif isinstance(node,CitationReference):parts.append((" " if previous_citation else "&nbsp;")+f'<span class="citation">[{node.reference_id}]</span>');previous_citation=True
                elif isinstance(node,SoftBreak):parts.append("<span class=\"break\"> </span>");previous_citation=False
                elif isinstance(node,Bold):parts.append(f"<strong>{inline_html(node.children)}</strong>");previous_citation=False
                elif isinstance(node,Italic):parts.append(f"<em>{inline_html(node.children)}</em>");previous_citation=False
                elif isinstance(node,Link):parts.append(f'<a href="{escape(node.href,quote=True)}">{inline_html(node.children)}</a>');previous_citation=False
            return "".join(parts)
        parts=[f"<h1>{safe_text(document.title)}</h1>"]
        if document.context_metadata.get("export_options",{}).get("include_generated_timestamp",True): parts.append(f'<p class="meta">DOCX Preview • Generated {document.generated_at:%d %b %Y, %H:%M UTC}</p>')
        if document.query: parts += ["<h2>Question</h2>", f"<p>{safe_text(document.query)}</p>"]
        for block in document.blocks:
            if isinstance(block, HeadingBlock): parts.append(f"<h{min(block.level,3)}>{safe_text(block.text)}</h{min(block.level,3)}>")
            elif isinstance(block, ParagraphBlock): parts.append(f"<p>{inline_html(block.inline_nodes)}</p>")
            elif isinstance(block, (BulletListBlock, NumberedListBlock)):
                tag="ul" if isinstance(block,BulletListBlock) else "ol";items=[];has_tasks=isinstance(block,BulletListBlock) and any(state is not None for state in block.task_states)
                for index,item in enumerate(block.items):
                    content=" ".join(inline_html(p.inline_nodes) for p in item if isinstance(p,ParagraphBlock));state=block.task_states[index] if isinstance(block,BulletListBlock) and index<len(block.task_states) else None
                    if state is None and has_tasks:items.append(f'<li class="task-item"><span class="task-marker">•</span><span class="task-content">{content}</span></li>')
                    elif state is None:items.append(f"<li>{content}</li>")
                    else:items.append(f'<li class="task-item"><span class="task-marker">{"☒" if state else "☐"}</span><span class="task-content">{content}</span></li>')
                list_class=' class="task-list"' if has_tasks else "";parts.append(f"<{tag}{list_class}>"+"".join(items)+f"</{tag}>")
            elif isinstance(block, TableBlock): parts.append('<table class="export-table"><thead><tr>'+"".join(f"<th>{safe_text(x)}</th>" for x in block.headers)+"</tr></thead><tbody>"+"".join("<tr>"+"".join(f"<td>{safe_text(x)}</td>" for x in row)+"</tr>" for row in block.rows)+"</tbody></table>")
            elif isinstance(block, CodeBlock): parts.append(f"<pre><code>{safe_text(block.code)}</code></pre>")
            elif isinstance(block, QuoteBlock): parts.append("<blockquote>"+"".join(f"<p>{inline_html(p.inline_nodes)}</p>" for p in block.blocks if isinstance(p,ParagraphBlock))+"</blockquote>")
            elif isinstance(block, HorizontalRuleBlock): parts.append("<hr>")
        if document.sources: parts += ["<h2>Sources</h2>", "<ul>"+"".join(f"<li>[{s.citation_number}] {safe_text(s.document_title)}"+(f", p. {s.page_number}" if s.page_number else "")+"</li>" for s in document.sources)+"</ul>"]
        clean=bleach.clean('<div class="page">'+"".join(parts)+"</div>", tags={"div","h1","h2","h3","p","ul","ol","li","table","thead","tbody","tr","th","td","pre","code","blockquote","strong","em","a","span","hr"}, attributes={"div":["class"],"p":["class"],"a":["href","title"],"span":["class"],"ul":["class"],"li":["class"],"code":["class"],"table":["class"],"th":["colspan","rowspan"],"td":["colspan","rowspan"]}, protocols={"http","https","mailto"}, strip=True)
        css="html,body{margin:0;min-height:100%;background:#eef1ed}body{padding:24px;font:16px/1.52 Arial,sans-serif;color:#172018}.page{box-sizing:border-box;max-width:760px;margin:0 auto;padding:46px 54px;background:#fff;border:1px solid #dfe5dc;box-shadow:0 4px 16px rgba(24,38,22,.08)}h1,h2,h3{color:#245c20;line-height:1.25;margin:0}h1{font-size:24px;margin-top:18px;margin-bottom:7px}h2{font-size:19px;margin-top:16px;margin-bottom:6px}h3{font-size:16px;margin-top:13px;margin-bottom:5px}.page>h1:first-child{font-size:32px;line-height:1.18;margin:0 0 5px}.meta{margin:0 0 18px;color:#737d72;font-size:12px;line-height:1.35}p{margin:0 0 8px}ul,ol{margin:4px 0 10px;padding-left:25px}li{margin:2px 0;padding-left:2px}li>p:first-child{margin-top:0}li>p:last-child{margin-bottom:0}.task-list{list-style:none;padding-left:0;margin:5px 0 11px}.task-item{display:grid;grid-template-columns:22px minmax(0,1fr);column-gap:5px;margin:3px 0;padding:0}.task-marker{grid-column:1;line-height:1.52;color:#245c20}.task-content{grid-column:2;min-width:0}.citation{white-space:nowrap;color:#355f31;font-size:.9em}table{border-collapse:collapse;width:100%;margin:8px 0 12px}th,td{border:1px solid #bcc8b8;padding:7px;text-align:left;vertical-align:top}th{background:#e8f1e5}pre{margin:8px 0 12px;background:#f3f5f2;padding:10px;overflow:auto;font-size:13px;line-height:1.45}blockquote{margin:8px 0 12px;border-left:3px solid #7ca475;padding:5px 0 5px 12px;color:#394638}hr{border:0;border-top:1px solid #cbd5c8;margin:14px 0}@media(max-width:640px){body{padding:0}.page{border:0;box-shadow:none;padding:28px 22px}}"
        return f"<!doctype html><html><head><meta charset=\"utf-8\"><title>DOCX Preview</title><style>{css}</style></head><body>{clean}</body></html>"
