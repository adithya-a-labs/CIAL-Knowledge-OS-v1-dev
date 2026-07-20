"""Export artifact discovery service."""

from __future__ import annotations

from datetime import datetime, timezone
import hashlib
from pathlib import Path

from backend.app.core.paths import OUTPUTS_ROOT, REPO_ROOT
from backend.app.schemas.exports import ExportFile


class ExportService:
    def __init__(self, outputs_root: Path = OUTPUTS_ROOT) -> None:
        self.outputs_root = outputs_root

    def list_exports(self) -> list[ExportFile]:
        if not self.outputs_root.exists():
            return []
        allowed = {".csv", ".xlsx", ".html", ".json", ".jsonl", ".log", ".txt"}
        files = [
            path for path in self.outputs_root.rglob("*")
            if path.is_file() and path.suffix.casefold() in allowed
        ]
        exports: list[ExportFile] = []
        for path in sorted(files, key=lambda item: item.stat().st_mtime, reverse=True):
            stat = path.stat()
            relative = path.relative_to(REPO_ROOT).as_posix()
            exports.append(
                ExportFile(
                    id=hashlib.sha1(relative.encode("utf-8")).hexdigest()[:16],
                    name=path.name,
                    path=relative,
                    type=path.suffix.casefold().lstrip(".") or "file",
                    size_bytes=stat.st_size,
                    modified_at=datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc).isoformat(),
                )
            )
        return exports[:200]

    def export_chat_message(self, message, question: str | None, format: str, include_sources: bool, include_metadata: bool) -> tuple[str, str]:
        from docx import Document as WordDocument
        export_dir = self.outputs_root / "chat"
        export_dir.mkdir(parents=True, exist_ok=True)
        stamp = datetime.now(timezone.utc).strftime("%Y%m%d-%H%M%S")
        filename = f"cial-response-{message.id}-{stamp}.{format}"
        path = export_dir / filename
        lines = ["CIAL Knowledge OS", "", "Question", question or "Unavailable", "", "Answer", message.content]
        if include_sources and message.sources:
            lines += ["", "Sources"] + [f"- {s.get('document_name', 'Unknown')}" + (f", p. {s.get('page_number') or s.get('page')}" if s.get('page_number') or s.get('page') else "") for s in message.sources]
        if include_metadata:
            lines += ["", "Generation metadata", str(message.metadata_ or {}), f"Exported: {datetime.now(timezone.utc).isoformat()}"]
        if format == "docx":
            doc = WordDocument(); doc.add_heading(lines[0], 0)
            for line in lines[1:]: doc.add_paragraph(line)
            doc.save(path)
        elif format == "pdf":
            try:
                from reportlab.lib.pagesizes import A4
                from reportlab.pdfgen import canvas
            except ImportError as exc:
                raise RuntimeError("PDF export support is unavailable.") from exc
            pdf = canvas.Canvas(str(path), pagesize=A4); _, height = A4; y = height - 54
            for line in lines:
                for part in ([line[i:i+95] for i in range(0, len(line), 95)] or [""]):
                    if y < 54: pdf.showPage(); y = height - 54
                    pdf.drawString(54, y, part); y -= 16
            pdf.save()
        else:
            raise ValueError("Unsupported export format.")
        return filename, f"/api/exports/chat/{filename}"
